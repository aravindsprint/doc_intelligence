import frappe
from frappe.model.document import Document
import os, json


class AIDocument(Document):
    def before_save(self):
        if not self.title and self.file_attachment:
            self.title = os.path.splitext(os.path.basename(self.file_attachment))[0].replace("-", " ").replace("_", " ").title()

    def after_insert(self):
        if self.file_attachment and self.status == "Pending":
            frappe.enqueue(
                "doc_intelligence.doc_intelligence.doctype.ai_document.ai_document.process_document",
                doc_name=self.name, queue="long", timeout=300,
            )

    def on_update(self):
        if self.status == "Pending" and self.file_attachment and not self.is_new():
            frappe.enqueue(
                "doc_intelligence.doc_intelligence.doctype.ai_document.ai_document.process_document",
                doc_name=self.name, queue="long", timeout=300,
            )


def fail_stuck_processing_documents():
    """
    Safety net: if a document has been stuck in Processing for longer than
    any single processing attempt could reasonably take (well past the
    background job's own 300s timeout), the worker that was handling it
    almost certainly died or was killed mid-request without ever reaching
    the except block below to mark it Failed. Runs every 15 minutes (see
    hooks.py) and cleans up anything left stranded like that.
    """
    cutoff = frappe.utils.add_to_date(frappe.utils.now_datetime(), minutes=-15)
    stuck = frappe.get_all(
        "AI Document",
        filters={"status": "Processing", "modified": ["<", cutoff]},
        pluck="name",
    )
    for doc_name in stuck:
        frappe.db.set_value(
            "AI Document", doc_name, "status", "Failed",
            update_modified=True,
        )
        frappe.log_error(
            f"AI Document {doc_name} was stuck in Processing for over 15 minutes "
            f"and was auto-marked Failed by the cleanup job.",
            "AI Document stuck-processing cleanup",
        )
    if stuck:
        frappe.db.commit()  # nosemgrep: frappe-manual-commit -- scheduled job, no request-level auto-commit to rely on


def _as_text(value):
    """The LLM is asked for 'summary'/'entities' as plain strings, but it
    sometimes returns a JSON array instead (e.g. entities as a list of
    strings) despite the prompt. Frappe's Long Text fields can't store a
    list, so coerce anything non-string into readable text before it ever
    reaches doc.save()."""
    if isinstance(value, str):
        return value
    if isinstance(value, (list, tuple)):
        return "\n".join(f"- {_as_text(v)}" for v in value)
    if isinstance(value, dict):
        return "\n".join(f"- {k}: {_as_text(v)}" for k, v in value.items())
    if value is None:
        return ""
    return str(value)


def process_document(doc_name):
    doc = frappe.get_doc("AI Document", doc_name)
    try:
        doc.status = "Processing"
        doc.save(ignore_permissions=True)
        frappe.db.commit()  # nosemgrep: frappe-manual-commit -- must be visible before the long-running LLM call, in case the job is killed mid-flight

        raw_text = extract_text(doc.file_attachment)
        doc.raw_text = raw_text

        settings = frappe.get_single("Doc Intelligence Settings")
        from doc_intelligence.doc_intelligence.llm_engine import analyse_document
        result = analyse_document(raw_text, doc.document_type or "Document", None, settings.max_tokens_per_request or 2000)

        doc.summary = _as_text(result.get("summary", ""))
        doc.key_entities = _as_text(result.get("entities", ""))
        tables = result.get("tables", [])
        doc.extracted_table = json.dumps(tables) if tables else ""
        meta = result.get("_meta", {})
        doc.token_count = meta.get("tokens_out", 0)
        doc.provider_used = meta.get("provider", "")
        if meta.get("fallback_used"):
            doc.provider_used = f"{meta.get('provider')} (fallback)"
        doc.status = "Ready"
        doc.processed_on = frappe.utils.now_datetime()
        doc.save(ignore_permissions=True)
        frappe.db.commit()  # nosemgrep: frappe-manual-commit -- background job, commits its own result explicitly

    except Exception:
        frappe.log_error(frappe.get_traceback(), f"AI Document processing failed: {doc_name}")
        # This save must never itself be allowed to fail silently — if it
        # does (a validation error, a timestamp race with another worker,
        # anything), the document is left permanently stuck instead of
        # cleanly marked Failed. Fall back to a raw, hook-bypassing update
        # as a last resort so the status change always lands no matter what.
        try:
            doc.reload()
            doc.status = "Failed"
            doc.save(ignore_permissions=True)
            frappe.db.commit()  # nosemgrep: frappe-manual-commit -- failure path must land even if the rest of the job never committed
        except Exception:
            frappe.log_error(frappe.get_traceback(), f"AI Document failure-handling itself failed: {doc_name}")
            frappe.db.set_value("AI Document", doc_name, "status", "Failed", update_modified=True)
            frappe.db.commit()  # nosemgrep: frappe-manual-commit -- last-resort fallback, must guarantee the status change lands


def extract_text(file_url):
    site_path = frappe.get_site_path()
    if file_url.startswith("/private/files/"):
        file_path = os.path.join(site_path, "private", "files", os.path.basename(file_url))
    else:
        file_path = os.path.join(site_path, "public", "files", os.path.basename(file_url))

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == ".docx":
        from docx import Document
        d = Document(file_path)
        return "\n".join(p.text for p in d.paragraphs)
    elif ext in (".jpg", ".jpeg", ".png", ".webp", ".gif"):
        from doc_intelligence.doc_intelligence.llm_engine import vision_extract_text
        return vision_extract_text(file_path)
    else:
        frappe.throw(f"Unsupported file type: {ext}. Supported: PDF, DOCX, and images (JPG/PNG/WEBP).")
